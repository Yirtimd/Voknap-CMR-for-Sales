import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile, is_zipfile

import fitz
import pytesseract
from docx import Document
from PIL import Image
from pypdf import PdfReader

from app.core.config import settings


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MIME_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
}
ALLOWED_CONTENT_TYPES = {
    ".pdf": {"application/pdf", "application/octet-stream"},
    ".docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/zip",
        "application/octet-stream",
    },
    ".txt": {"text/plain", "application/octet-stream"},
}


class KnowledgeFileError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedKnowledgeFile:
    filename: str
    extension: str
    mime_type: str
    text: str
    extraction_method: str
    source_pages: int | None = None


def parse_knowledge_file(filename: str, content_type: str | None, data: bytes) -> ParsedKnowledgeFile:
    safe_name = sanitize_filename(filename)
    extension = Path(safe_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise KnowledgeFileError("Supported file types: PDF, DOCX, TXT")
    if not data:
        raise KnowledgeFileError("Uploaded file is empty")
    if len(data) > settings.knowledge_max_upload_bytes:
        raise KnowledgeFileError(
            f"File exceeds {settings.knowledge_max_upload_bytes // 1_000_000} MB limit"
        )
    normalized_content_type = (content_type or "").split(";", 1)[0].strip().lower()
    if normalized_content_type and normalized_content_type not in ALLOWED_CONTENT_TYPES[extension]:
        raise KnowledgeFileError("File content type does not match its extension")

    if extension == ".pdf":
        text, extraction_method, source_pages = _parse_pdf(data)
    elif extension == ".docx":
        text = _parse_docx(data)
        extraction_method = "docx"
        source_pages = None
    else:
        text = _parse_txt(data)
        extraction_method = "text"
        source_pages = None

    clean_text = text.strip()
    if len(clean_text) < 20:
        raise KnowledgeFileError("File contains too little extractable text")
    if len(clean_text) > settings.knowledge_max_extracted_chars:
        raise KnowledgeFileError(
            f"Extracted text exceeds {settings.knowledge_max_extracted_chars} character limit"
        )
    return ParsedKnowledgeFile(
        filename=safe_name,
        extension=extension.removeprefix("."),
        mime_type=MIME_TYPES[extension],
        text=clean_text,
        extraction_method=extraction_method,
        source_pages=source_pages,
    )


def sanitize_filename(filename: str) -> str:
    name = Path(filename or "document").name.strip()
    name = re.sub(r"[^\w.()\- ]+", "_", name, flags=re.UNICODE)
    name = re.sub(r"\s+", " ", name).strip(" .")
    if not name:
        raise KnowledgeFileError("File name is empty")
    stem = Path(name).stem[:180].rstrip(" .") or "document"
    suffix = Path(name).suffix.lower()[:10]
    return f"{stem}{suffix}"


def _parse_pdf(data: bytes) -> tuple[str, str, int]:
    if not data.startswith(b"%PDF-"):
        raise KnowledgeFileError("File extension is PDF, but content is not a PDF")
    try:
        reader = PdfReader(BytesIO(data), strict=False)
        if reader.is_encrypted and reader.decrypt("") == 0:
            raise KnowledgeFileError("Encrypted PDF files are not supported")
        if len(reader.pages) > settings.knowledge_max_pdf_pages:
            raise KnowledgeFileError(
                f"PDF exceeds {settings.knowledge_max_pdf_pages} page limit"
            )
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except KnowledgeFileError:
        raise
    except Exception as error:
        raise KnowledgeFileError("Cannot parse PDF file") from error

    ocr_page_indexes = [index for index, text in enumerate(pages) if len(text) < 20]
    if ocr_page_indexes and settings.knowledge_ocr_enabled:
        if len(ocr_page_indexes) > settings.knowledge_ocr_max_pages:
            raise KnowledgeFileError(
                f"PDF requires OCR for {len(ocr_page_indexes)} pages; "
                f"limit is {settings.knowledge_ocr_max_pages}"
            )
        try:
            ocr_text = _ocr_pdf_pages(data, ocr_page_indexes)
        except KnowledgeFileError:
            # A text PDF can legitimately contain blank separator pages. Do not
            # reject all useful text when OCR is unavailable on a local host.
            if len("\n".join(pages).strip()) < 20:
                raise
            return "\n\n".join(page for page in pages if page), "pypdf", len(pages)
        for page_index, text in ocr_text.items():
            pages[page_index] = text
        extraction_method = "pypdf+ocr"
    else:
        extraction_method = "pypdf"
    return "\n\n".join(page for page in pages if page), extraction_method, len(pages)


def _ocr_pdf_pages(data: bytes, page_indexes: list[int]) -> dict[int, str]:
    try:
        document = fitz.open(stream=data, filetype="pdf")
    except Exception as error:
        raise KnowledgeFileError("Cannot render PDF for OCR") from error
    extracted: dict[int, str] = {}
    try:
        for page_index in page_indexes:
            page = document.load_page(page_index)
            pixmap = page.get_pixmap(dpi=settings.knowledge_ocr_dpi, alpha=False)
            image = Image.open(BytesIO(pixmap.tobytes("png")))
            try:
                text = pytesseract.image_to_string(
                    image,
                    lang=settings.knowledge_ocr_languages,
                    timeout=settings.knowledge_ocr_page_timeout_seconds,
                )
            except pytesseract.TesseractNotFoundError as error:
                raise KnowledgeFileError(
                    "OCR engine is unavailable; install Tesseract and configured language packs"
                ) from error
            except (pytesseract.TesseractError, RuntimeError) as error:
                raise KnowledgeFileError(f"OCR failed on PDF page {page_index + 1}") from error
            extracted[page_index] = text.strip()
    finally:
        document.close()
    return extracted


def _parse_docx(data: bytes) -> str:
    buffer = BytesIO(data)
    if not is_zipfile(buffer):
        raise KnowledgeFileError("File extension is DOCX, but content is not a DOCX archive")
    buffer.seek(0)
    try:
        with ZipFile(buffer) as archive:
            infos = archive.infolist()
            if "word/document.xml" not in {item.filename for item in infos}:
                raise KnowledgeFileError("DOCX document body is missing")
            if len(infos) > 2_000:
                raise KnowledgeFileError("DOCX archive contains too many entries")
            if sum(item.file_size for item in infos) > 50_000_000:
                raise KnowledgeFileError("DOCX expanded content exceeds 50 MB limit")
            if any(item.flag_bits & 0x1 for item in infos):
                raise KnowledgeFileError("Encrypted DOCX files are not supported")
        buffer.seek(0)
        document = Document(buffer)
    except KnowledgeFileError:
        raise
    except (BadZipFile, Exception) as error:
        raise KnowledgeFileError("Cannot parse DOCX file") from error

    parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _parse_txt(data: bytes) -> str:
    encodings = ["utf-8-sig"]
    if data.startswith((b"\xff\xfe", b"\xfe\xff")):
        encodings.append("utf-16")
    encodings.append("cp1251")
    for encoding in encodings:
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise KnowledgeFileError("TXT encoding is not supported; use UTF-8, UTF-16, or CP1251")
