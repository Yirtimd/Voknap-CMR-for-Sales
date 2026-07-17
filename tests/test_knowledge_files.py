from io import BytesIO
from uuid import uuid4

import pytest
import fitz
import pytesseract
from docx import Document

from app.modules.knowledge.files import (
    KnowledgeFileError,
    parse_knowledge_file,
)
from app.modules.knowledge.storage import resolve_local_storage_key, store_knowledge_file


def test_parse_utf8_and_cp1251_txt():
    utf8 = parse_knowledge_file(
        "playbook.txt",
        "text/plain; charset=utf-8",
        "Следующий шаг по сделке — отправить коммерческое предложение.".encode(),
    )
    cp1251 = parse_knowledge_file(
        "client.txt",
        "text/plain",
        "Клиент ожидает встречу с менеджером на следующей неделе.".encode("cp1251"),
    )

    assert "коммерческое предложение" in utf8.text
    assert "следующей неделе" in cp1251.text
    assert utf8.mime_type == "text/plain"


def test_parse_docx_paragraphs_and_tables():
    document = Document()
    document.add_paragraph("План внедрения CRM и обучения менеджеров.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Следующий шаг"
    table.cell(0, 1).text = "Согласовать пилот"
    buffer = BytesIO()
    document.save(buffer)

    parsed = parse_knowledge_file(
        "implementation.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buffer.getvalue(),
    )

    assert "План внедрения CRM" in parsed.text
    assert "Следующий шаг | Согласовать пилот" in parsed.text


def test_parse_pdf_text():
    parsed = parse_knowledge_file("proposal.pdf", "application/pdf", _minimal_text_pdf())

    assert "Hello PDF knowledge base" in parsed.text
    assert parsed.mime_type == "application/pdf"
    assert parsed.extraction_method == "pypdf"
    assert parsed.source_pages == 1


def test_scanned_pdf_uses_ocr(monkeypatch):
    document = fitz.open()
    document.new_page()
    data = document.tobytes()
    document.close()
    monkeypatch.setattr(
        "app.modules.knowledge.files.pytesseract.image_to_string",
        lambda *_args, **_kwargs: "Распознанный текст сканированного коммерческого предложения.",
    )

    parsed = parse_knowledge_file("scan.pdf", "application/pdf", data)

    assert "Распознанный текст" in parsed.text
    assert parsed.extraction_method == "pypdf+ocr"


def test_text_pdf_with_blank_page_survives_missing_local_ocr(monkeypatch):
    document = fitz.open()
    text_page = document.new_page()
    text_page.insert_text(
        (72, 72),
        "Useful text from a normal PDF page for company knowledge.",
    )
    document.new_page()
    data = document.tobytes()
    document.close()

    def missing_tesseract(*_args, **_kwargs):
        raise pytesseract.TesseractNotFoundError()

    monkeypatch.setattr(
        "app.modules.knowledge.files.pytesseract.image_to_string",
        missing_tesseract,
    )

    parsed = parse_knowledge_file("mixed.pdf", "application/pdf", data)

    assert "Useful text" in parsed.text
    assert parsed.extraction_method == "pypdf"
    assert parsed.source_pages == 2


def test_reject_mismatched_or_unsupported_files():
    with pytest.raises(KnowledgeFileError, match="content type"):
        parse_knowledge_file("fake.pdf", "text/plain", b"%PDF-fake content with enough text")
    with pytest.raises(KnowledgeFileError, match="Supported file types"):
        parse_knowledge_file("sheet.xlsx", "application/octet-stream", b"x" * 100)


def test_storage_uses_tenant_directory_and_blocks_traversal(tmp_path, monkeypatch):
    monkeypatch.setattr("app.modules.knowledge.storage.settings.knowledge_storage_backend", "local")
    monkeypatch.setattr("app.modules.knowledge.storage.settings.knowledge_upload_dir", str(tmp_path))
    tenant_id = uuid4()

    backend, storage_key = store_knowledge_file(
        tenant_id,
        "../КП клиента.txt",
        "Документ компании с коммерческими условиями.".encode(),
        "text/plain",
    )
    path = resolve_local_storage_key(storage_key)

    assert backend == "local"
    assert path.read_text() == "Документ компании с коммерческими условиями."
    assert storage_key.startswith(str(tenant_id))
    assert path.is_relative_to(tmp_path)
    with pytest.raises(KnowledgeFileError, match="Invalid knowledge file storage path"):
        resolve_local_storage_key("../../outside.txt")


def _minimal_text_pdf() -> bytes:
    stream = b"BT /F1 12 Tf 72 720 Td (Hello PDF knowledge base) Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    result = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(result))
        result.extend(f"{index} 0 obj\n".encode())
        result.extend(body)
        result.extend(b"\nendobj\n")
    xref_offset = len(result)
    result.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    result.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        result.extend(f"{offset:010d} 00000 n \n".encode())
    result.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode()
    )
    return bytes(result)
